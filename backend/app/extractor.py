"""Text extraction from PDF / DOCX / TXT / MD, plus cleaning/preprocessing."""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path


def _extract_pdf_bytes(data: bytes) -> str:
    parts: list[str] = []
    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
    except Exception:  # noqa: BLE001  -> fall back to pypdf/PyPDF2
        try:
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader
            reader = PdfReader(BytesIO(data))
            parts = [(p.extract_text() or "") for p in reader.pages]
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"Failed to read PDF: {exc}") from exc
    return "\n".join(parts)


def _extract_docx_bytes(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    paras = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.append(cell.text)
    return "\n".join(paras)


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_bytes(data)
    if ext == ".docx":
        return _extract_docx_bytes(data)
    if ext in (".txt", ".md", ".markdown"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {ext}")


def extract_text(path: Path) -> str:
    return extract_text_from_bytes(path.name, Path(path).read_bytes())


def clean_text(text: str) -> str:
    text = (text or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    return text.strip()


def normalize_for_matching(text: str) -> str:
    text = (text or "").lower()
    text = re.sub(r"[^a-z0-9+#./\s-]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()
